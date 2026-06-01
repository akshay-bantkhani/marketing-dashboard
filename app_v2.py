import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime, timedelta, date
import json
import os
import pickle
import numpy as np
import requests

st.set_page_config(page_title="Contify Marketing Strategy Dashboard", page_icon="🚀", layout="wide")

# =====================================================
# PATHS & CONFIG
# =====================================================
APP_DIR = os.path.dirname(os.path.abspath(__file__))
TOKEN_FILE = os.path.join(APP_DIR, "token.pickle")
CONFIG_FILE = os.path.join(APP_DIR, "config.json")
SERVICE_ACCOUNT_FILE = os.path.join(os.path.expanduser("~"), "Downloads", "seo-report-weekly-054644ce3e32.json")
OAUTH_CLIENT_FILE = os.path.join(os.path.expanduser("~"), "Downloads", "client_secret_903915771943-pngfdafiam32olbubqq8kiqb60lh9f73.apps.googleusercontent.com.json")

SCOPES = ['https://www.googleapis.com/auth/analytics.readonly', 'https://www.googleapis.com/auth/webmasters.readonly']
DEFAULT_CONFIG = {"ga4_property": "250811000", "gsc_site": "https://www.contify.com/", "auth_method": "demo",
                  "hubspot_token": "",
                  "targets": {"overall_traffic": 5538, "organic_traffic": 1154, "overall_users": 4562, "organic_users": 706, "mqls": 13, "leads": 50, "sqls": 6, "pipeline": 50000}}

AI_SOURCES = ['chatgpt.com', 'chat.openai.com', 'gemini.google.com', 'claude.ai', 'perplexity.ai',
              'copilot.microsoft.com', 'you.com', 'bard.google.com', 'poe.com', 'phind.com']

# =====================================================
# CSS — Strategy Dashboard Style
# =====================================================
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');

html, body, [data-testid="stApp"], [data-testid="stApp"] > div { background-color: #f8f9fc !important; color: #1e293b !important; }
html, body, [class*="css"], [data-testid="stApp"] * { font-family: 'Inter', sans-serif; }
.main .block-container { max-width: 1200px; padding: 1.5rem 2rem; }

[data-testid="stSidebar"] { background-color: #ffffff !important; border-right: 1px solid #e5e7eb !important; }
[data-testid="stSidebar"] * { color: #1e293b !important; }
[data-testid="stSidebar"] [data-testid="stMarkdownContainer"] h2,
[data-testid="stSidebar"] [data-testid="stMarkdownContainer"] strong { color: #0f3460 !important; }

[data-testid="stNumberInput"] input, [data-testid="stTextInput"] input { background: #f8fafc !important; border: 1px solid #d1d5db !important; border-radius: 8px !important; color: #1e293b !important; }
button[kind="secondary"], [data-testid="stButton"] button {
    background: linear-gradient(135deg, #0f3460, #1a56db) !important; color: #ffffff !important;
    border: none !important; border-radius: 8px !important; font-weight: 600 !important; padding: 8px 20px !important;
}
button[kind="secondary"]:hover, [data-testid="stButton"] button:hover { opacity: 0.9 !important; }

/* Dashboard Header */
.dash-header {
    background: linear-gradient(135deg, #0f172a 0%, #1e40af 50%, #3b82f6 100%);
    padding: 28px 36px; border-radius: 16px; color: white; margin-bottom: 24px;
    position: relative; overflow: hidden;
}
.dash-header::after { content: ''; position: absolute; right: -40px; top: -40px; width: 200px; height: 200px;
    background: rgba(255,255,255,0.05); border-radius: 50%; }
.dash-header h1 { margin: 0; font-size: 1.7rem; font-weight: 800; letter-spacing: -0.5px; color: #ffffff !important; }
.dash-header p { margin: 6px 0 0; opacity: 0.8; font-size: 0.88rem; color: #ffffff !important; }

/* Strategy Section */
.strat-section {
    background: #ffffff; border-radius: 14px; padding: 0; margin-bottom: 20px;
    box-shadow: 0 1px 4px rgba(0,0,0,0.06); border: 1px solid #e5e7eb; overflow: hidden;
}
.strat-header {
    padding: 16px 24px; font-size: 1rem; font-weight: 700; color: #0f172a;
    border-bottom: 1px solid #e5e7eb; display: flex; align-items: center; gap: 10px;
    background: #fafbfd;
}
.strat-body { padding: 20px 24px; }

/* Funnel Cards */
.funnel-row { display: flex; gap: 0; align-items: stretch; margin-bottom: 20px; }
.funnel-card {
    flex: 1; text-align: center; padding: 22px 10px; position: relative; min-width: 0;
    border: 1px solid #e5e7eb; background: #fff;
}
.funnel-card:first-child { border-radius: 12px 0 0 12px; }
.funnel-card:last-child { border-radius: 0 12px 12px 0; }
.funnel-card .f-label { font-size: 0.68rem; color: #6b7280; text-transform: uppercase; letter-spacing: 0.7px; font-weight: 700; margin-bottom: 4px; }
.funnel-card .f-val { font-size: 1.6rem; font-weight: 800; color: #0f172a; }
.funnel-card .f-change { font-size: 0.78rem; font-weight: 600; margin-top: 4px; }
.funnel-card .f-arrow { position: absolute; right: -10px; top: 50%; transform: translateY(-50%); z-index: 2;
    width: 20px; height: 20px; background: #e5e7eb; border-radius: 50%; display: flex; align-items: center; justify-content: center;
    font-size: 0.7rem; color: #6b7280; }
.funnel-card .f-conv { font-size: 0.65rem; color: #1a56db; font-weight: 600; margin-top: 2px; }

/* Metric Cards */
.m-card {
    background: #ffffff; border: 1px solid #e5e7eb; border-radius: 12px; padding: 18px 16px;
    text-align: center; box-shadow: 0 1px 3px rgba(0,0,0,0.04);
}
.m-card .m-label { font-size: 0.7rem; color: #6b7280; text-transform: uppercase; letter-spacing: 0.6px; font-weight: 700; margin-bottom: 6px; }
.m-card .m-val { font-size: 1.5rem; font-weight: 800; color: #0f172a; }
.m-card .m-sub { font-size: 0.78rem; font-weight: 600; margin-top: 4px; }

/* Insight Box — Strategy Style */
.strategy-box {
    background: #f0fdf4; border-left: 4px solid #10b981; padding: 16px 20px;
    border-radius: 0 10px 10px 0; margin: 14px 0; font-size: 0.88rem; line-height: 1.7;
}
.strategy-box.warning { background: #fffbeb; border-left-color: #f59e0b; }
.strategy-box.action { background: #eff6ff; border-left-color: #1a56db; }
.strategy-box strong { color: #0f172a; }
.strategy-box ul { margin: 6px 0 0; padding-left: 18px; }
.strategy-box li { margin-bottom: 4px; }

/* Data Tables */
.data-table { width: 100%; border-collapse: separate; border-spacing: 0; border-radius: 10px; overflow: hidden; border: 1px solid #e5e7eb; font-size: 0.88rem; }
.data-table th { background: #f8fafc; color: #374151; padding: 11px 16px; font-weight: 700; text-align: left; border-bottom: 2px solid #e5e7eb; font-size: 0.78rem; text-transform: uppercase; letter-spacing: 0.3px; }
.data-table td { padding: 10px 16px; border-bottom: 1px solid #f0f2f5; color: #1e293b; }
.data-table tr:last-child td { border-bottom: none; }
.data-table tr:hover td { background: #f8fafc; }

/* Detail Table */
.detail-table { width: 100%; border-collapse: separate; border-spacing: 0; border-radius: 8px; overflow: hidden; border: 1px solid #dbeafe; font-size: 0.82rem; margin-top: 8px; }
.detail-table th { background: #eff6ff; color: #1e40af; padding: 9px 12px; font-weight: 600; text-align: left; font-size: 0.75rem; text-transform: uppercase; letter-spacing: 0.3px; }
.detail-table td { padding: 8px 12px; border-bottom: 1px solid #f0f7ff; color: #374151; }
.detail-table tr:last-child td { border-bottom: none; }

/* Tags */
.tag { padding: 3px 10px; border-radius: 6px; font-size: 0.65rem; font-weight: 700; letter-spacing: 0.3px; }
.tag-green { background: #dcfce7; color: #166534; }
.tag-yellow { background: #fef3c7; color: #92400e; }
.tag-red { background: #fee2e2; color: #991b1b; }
.tag-blue { background: #dbeafe; color: #1e40af; }
.tag-current { background: #1a56db; color: white; padding: 3px 10px; border-radius: 6px; font-size: 0.7rem; font-weight: 700; }

.pos { color: #059669; font-weight: 700; }
.neg { color: #dc2626; font-weight: 700; }

/* Charts */
[data-testid="stPlotlyChart"] { border-radius: 12px; overflow: hidden; border: 1px solid #e5e7eb; background: #ffffff; padding: 6px; margin-bottom: 12px; }
[data-testid="stExpander"] { border: 1px solid #e5e7eb !important; border-radius: 10px !important; margin: 8px 0 !important; background: #ffffff !important; }
[data-testid="stExpander"] summary { font-weight: 600 !important; color: #1a56db !important; }

#MainMenu, footer, header { visibility: hidden; }
</style>
""", unsafe_allow_html=True)

# =====================================================
# CONFIG HELPERS
# =====================================================
def load_config():
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, 'r') as f:
            return {**DEFAULT_CONFIG, **json.load(f)}
    return DEFAULT_CONFIG.copy()

def save_config(c):
    with open(CONFIG_FILE, 'w') as f:
        json.dump(c, f, indent=2)

# =====================================================
# WEEK HELPERS
# =====================================================
def get_last_n_weeks(n=5, ref_date=None):
    today = ref_date or datetime.now().date()
    # Find the most recent completed Sunday (end of last full Mon–Sun week)
    days_since_sunday = (today.weekday() + 1) % 7  # 0 if today is Sunday
    last_sunday = today - timedelta(days=days_since_sunday)
    if last_sunday == today:
        last_sunday -= timedelta(days=7)
    weeks = []
    for i in range(n):
        sun = last_sunday - timedelta(weeks=i)
        mon = sun - timedelta(days=6)  # Monday = Sunday - 6
        weeks.append((mon, sun, f"{mon.strftime('%d %b')} – {sun.strftime('%d %b')}"))
    weeks.reverse()
    return weeks

def assign_weeks(df, weeks, date_col='Date', date_fmt='%Y%m%d'):
    df = df.copy()
    df['_date'] = pd.to_datetime(df[date_col], format=date_fmt, errors='coerce')
    df['Week'] = None; df['Week_Label'] = None; df['Week_Idx'] = -1
    for i, (m, s, lbl) in enumerate(weeks):
        mask = (df['_date'].dt.date >= m) & (df['_date'].dt.date <= s)
        df.loc[mask, 'Week'] = i; df.loc[mask, 'Week_Label'] = lbl; df.loc[mask, 'Week_Idx'] = i
    return df[df['Week'].notna()].copy()

# =====================================================
# AUTH
# =====================================================
def get_oauth_credentials():
    try:
        from google_auth_oauthlib.flow import InstalledAppFlow
        from google.auth.transport.requests import Request
    except ImportError:
        st.error("Missing: `pip install google-auth-oauthlib`"); return None
    creds = None
    if os.path.exists(TOKEN_FILE):
        with open(TOKEN_FILE, 'rb') as f: creds = pickle.load(f)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            try: creds.refresh(Request())
            except Exception: creds = None
        if not creds:
            if not os.path.exists(OAUTH_CLIENT_FILE):
                st.error(f"OAuth file missing: `{OAUTH_CLIENT_FILE}`"); return None
            flow = InstalledAppFlow.from_client_secrets_file(OAUTH_CLIENT_FILE, SCOPES)
            for port in [8090, 8091, 8092, 8093, 8094]:
                try: creds = flow.run_local_server(port=port, prompt='consent', access_type='offline'); break
                except OSError: continue
        with open(TOKEN_FILE, 'wb') as f: pickle.dump(creds, f)
    return creds

def _get_creds():
    if os.path.exists(TOKEN_FILE):
        with open(TOKEN_FILE, 'rb') as f: creds = pickle.load(f)
        if creds and creds.expired and creds.refresh_token:
            from google.auth.transport.requests import Request
            creds.refresh(Request())
            with open(TOKEN_FILE, 'wb') as f: pickle.dump(creds, f)
        return creds
    try:
        from google.oauth2 import service_account
        if os.path.exists(SERVICE_ACCOUNT_FILE):
            return service_account.Credentials.from_service_account_file(SERVICE_ACCOUNT_FILE, scopes=SCOPES)
    except Exception: pass
    return None

# =====================================================
# GA4 API
# =====================================================
def _ga4_report(creds, prop, start, end, dims, metrics, dim_filter=None, limit=10000):
    from google.analytics.data_v1beta import BetaAnalyticsDataClient
    from google.analytics.data_v1beta.types import RunReportRequest, DateRange, Dimension, Metric
    client = BetaAnalyticsDataClient(credentials=creds)
    req_kwargs = dict(property=f"properties/{prop}", date_ranges=[DateRange(start_date=start, end_date=end)],
                      dimensions=[Dimension(name=d) for d in dims], metrics=[Metric(name=m) for m in metrics], limit=limit)
    if dim_filter: req_kwargs['dimension_filter'] = dim_filter
    return client.run_report(RunReportRequest(**req_kwargs))

def _parse_response(response, dim_names, metric_names):
    rows = []
    for row in response.rows:
        r = {}
        for i, d in enumerate(dim_names): r[d] = row.dimension_values[i].value
        for i, m in enumerate(metric_names):
            v = row.metric_values[i].value
            r[m] = float(v) if '.' in v else int(v)
        rows.append(r)
    return pd.DataFrame(rows)

@st.cache_data(ttl=600, show_spinner=False)
def fetch_traffic(_ct, prop, start, end):
    creds = _get_creds()
    r = _ga4_report(creds, prop, start, end, ['date'], ['sessions','totalUsers','newUsers'])
    return _parse_response(r, ['Date'], ['Sessions','Users','New Users'])

@st.cache_data(ttl=600, show_spinner=False)
def fetch_channels(_ct, prop, start, end):
    creds = _get_creds()
    r = _ga4_report(creds, prop, start, end, ['sessionDefaultChannelGroup','date'], ['sessions','totalUsers','newUsers'])
    return _parse_response(r, ['Channel','Date'], ['Sessions','Users','New Users'])

@st.cache_data(ttl=600, show_spinner=False)
def fetch_landing_pages(_ct, prop, start, end):
    creds = _get_creds()
    r = _ga4_report(creds, prop, start, end, ['landingPage','date'], ['sessions','totalUsers','newUsers'])
    return _parse_response(r, ['Page','Date'], ['Sessions','Users','New Users'])

@st.cache_data(ttl=600, show_spinner=False)
def fetch_organic_traffic(_ct, prop, start, end):
    from google.analytics.data_v1beta.types import FilterExpression, Filter
    creds = _get_creds()
    filt = FilterExpression(filter=Filter(field_name="sessionDefaultChannelGroup", string_filter=Filter.StringFilter(value="Organic Search")))
    r = _ga4_report(creds, prop, start, end, ['date'], ['sessions','totalUsers','newUsers'], dim_filter=filt)
    return _parse_response(r, ['Date'], ['Sessions','Users','New Users'])

# =====================================================
# HUBSPOT API
# =====================================================
HUBSPOT_CONTACTS_URL = "https://api.hubapi.com/crm/v3/objects/contacts/search"
HUBSPOT_DEALS_URL = "https://api.hubapi.com/crm/v3/objects/deals/search"

def _hubspot_headers(token):
    return {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}

def _hubspot_search_paginated(url, token, payload):
    headers = _hubspot_headers(token)
    all_results = []
    after = None
    while True:
        body = payload.copy()
        if after: body["after"] = after
        try:
            resp = requests.post(url, headers=headers, json=body, timeout=30)
            resp.raise_for_status()
            data = resp.json()
        except requests.exceptions.RequestException as e:
            st.warning(f"HubSpot API error: {e}"); break
        results = data.get("results", [])
        all_results.extend(results)
        after = data.get("paging", {}).get("next", {}).get("after")
        if not after or not results: break
    return all_results

def _date_to_ms(d):
    return str(int(datetime.combine(d, datetime.min.time()).timestamp() * 1000))

@st.cache_data(ttl=600, show_spinner=False)
def fetch_hubspot_contacts(token, start_date, end_date):
    payload = {
        "filterGroups": [{"filters": [{"propertyName": "createdate", "operator": "BETWEEN",
            "highValue": _date_to_ms(end_date + timedelta(days=1)), "value": _date_to_ms(start_date)}]}],
        "properties": ["email","firstname","lastname","company","lifecyclestage","hs_analytics_source",
                        "hs_analytics_source_data_1","hs_analytics_source_data_2","industry","createdate",
                        "hs_lead_status","jobtitle","country"],
        "limit": 100
    }
    results = _hubspot_search_paginated(HUBSPOT_CONTACTS_URL, token, payload)
    rows = []
    for r in results:
        p = r.get("properties", {})
        rows.append({"email": p.get("email",""), "lifecyclestage": p.get("lifecyclestage",""),
            "source": p.get("hs_analytics_source",""), "source_detail": p.get("hs_analytics_source_data_1",""),
            "createdate": p.get("createdate",""), "name": f"{p.get('firstname','')} {p.get('lastname','')}".strip(),
            "company": p.get("company",""), "industry": p.get("industry",""), "jobtitle": p.get("jobtitle",""),
            "country": p.get("country",""), "lead_status": p.get("hs_lead_status","")})
    return pd.DataFrame(rows) if rows else pd.DataFrame(columns=["email","lifecyclestage","source","source_detail","createdate","name","company","industry","jobtitle","country","lead_status"])

def derive_mqls_from_contacts(contacts_df):
    """Derive MQLs from the contacts data — contacts whose lifecycle stage is MQL or higher."""
    if contacts_df.empty:
        return pd.DataFrame(columns=["email","source","createdate","name","company","industry","jobtitle","country"])
    # MQL = current stage is marketingqualifiedlead, salesqualifiedlead, or opportunity (they passed through MQL)
    mql_stages = ['marketingqualifiedlead', 'salesqualifiedlead', 'opportunity', 'customer', 'evangelist']
    mqls = contacts_df[contacts_df['lifecyclestage'].str.lower().isin(mql_stages)].copy()
    return mqls

def derive_sqls_from_contacts(contacts_df):
    """Derive SQLs from the contacts data — contacts whose lifecycle stage is SQL or higher."""
    if contacts_df.empty:
        return pd.DataFrame(columns=["email","source","createdate","name","company","industry","jobtitle","country"])
    sql_stages = ['salesqualifiedlead', 'opportunity', 'customer', 'evangelist']
    sqls = contacts_df[contacts_df['lifecyclestage'].str.lower().isin(sql_stages)].copy()
    return sqls

@st.cache_data(ttl=600, show_spinner=False)
def fetch_hubspot_deals(token, start_date, end_date):
    payload = {
        "filterGroups": [{"filters": [{"propertyName": "createdate", "operator": "BETWEEN",
            "highValue": _date_to_ms(end_date + timedelta(days=1)), "value": _date_to_ms(start_date)}]}],
        "properties": ["dealname","amount","dealstage","pipeline","closedate","createdate","hs_analytics_source"],
        "limit": 100
    }
    results = _hubspot_search_paginated(HUBSPOT_DEALS_URL, token, payload)
    rows = []
    for r in results:
        p = r.get("properties", {})
        amt = p.get("amount")
        rows.append({"dealname": p.get("dealname",""), "amount": float(amt) if amt else 0.0,
            "dealstage": p.get("dealstage",""), "pipeline": p.get("pipeline",""),
            "createdate": p.get("createdate",""), "closedate": p.get("closedate",""),
            "source": p.get("hs_analytics_source","")})
    return pd.DataFrame(rows) if rows else pd.DataFrame(columns=["dealname","amount","dealstage","pipeline","createdate","closedate","source"])

# =====================================================
# DEMO DATA
# =====================================================
def gen_demo_ga4(weeks):
    np.random.seed(42)
    all_dates = []
    for m, s, _ in weeks:
        all_dates.extend(pd.date_range(m, s))

    traffic_rows = []
    for d in all_dates:
        traffic_rows.append({'Date': d.strftime('%Y%m%d'), 'Sessions': np.random.randint(350, 500),
            'Users': np.random.randint(250, 400), 'New Users': np.random.randint(120, 280)})

    channels = ['Organic Search','Direct','Referral','Organic Social','Email','Paid Search']
    ch_rows = []
    for d in all_dates:
        for ch in channels:
            base = {'Organic Search': 140, 'Direct': 55, 'Referral': 25, 'Organic Social': 15, 'Email': 12, 'Paid Search': 18}.get(ch, 10)
            ch_rows.append({'Channel': ch, 'Date': d.strftime('%Y%m%d'),
                'Sessions': np.random.randint(max(1,base-15), base+20),
                'Users': np.random.randint(max(1,base-20), base+10),
                'New Users': np.random.randint(max(1,base//3), base//2+8)})

    pages = ['/','/blog/','/blog/competitive-intelligence-tools/','/platform/','/news-api/',
             '/blog/market-intelligence-vs-market-research/','/pricing/','/blog/swot-analysis/',
             '/case-studies/','/blog/competitor-analysis/','/resources/','/about/','/demo/',
             '/blog/sales-intelligence-tools/','/integrations/']
    lp_rows = []
    for d in all_dates:
        for p in pages:
            base = {'/': 40, '/blog/': 25, '/platform/': 15, '/pricing/': 12}.get(p, np.random.randint(3, 15))
            lp_rows.append({'Page': p, 'Date': d.strftime('%Y%m%d'),
                'Sessions': np.random.randint(max(1,base-5), base+8),
                'Users': np.random.randint(max(1,base-8), base+3),
                'New Users': np.random.randint(1, max(2,base//2))})

    org_rows = []
    for d in all_dates:
        org_rows.append({'Date': d.strftime('%Y%m%d'), 'Sessions': np.random.randint(110, 180),
            'Users': np.random.randint(85, 150), 'New Users': np.random.randint(55, 110)})

    return pd.DataFrame(traffic_rows), pd.DataFrame(ch_rows), pd.DataFrame(lp_rows), pd.DataFrame(org_rows)


def gen_demo_hubspot(weeks):
    np.random.seed(99)
    sources = ['ORGANIC_SEARCH','DIRECT_TRAFFIC','PAID_SEARCH','EMAIL_MARKETING','SOCIAL_MEDIA','REFERRALS','PAID_SOCIAL']
    industries = ['Technology','Financial Services','Healthcare','Manufacturing','Retail','Consulting','Media','Energy']
    companies = ['Acme Corp','TechVentures Inc','Global Finance Ltd','MedPharma Co','RetailMax','DataDriven LLC',
                 'CloudSync','InnovateTech','BrightPath','NexaGroup','Pinnacle Solutions','VelocityAI',
                 'Quantum Dynamics','AlphaStream','CoreLogic Systems','BlueSky Analytics','Vertex Partners',
                 'Prism Digital','Catalyst Labs','EmergePoint']
    first_names = ['James','Sarah','Michael','Emma','David','Lisa','Robert','Jennifer','William','Ashley',
                   'John','Maria','Daniel','Jessica','Mark','Laura','Steven','Rachel','Chris','Nicole']
    last_names = ['Johnson','Williams','Brown','Jones','Garcia','Miller','Davis','Rodriguez','Martinez','Anderson',
                  'Taylor','Thomas','Moore','Martin','Jackson','Thompson','White','Lopez','Lee','Harris']
    job_titles = ['VP Marketing','Director of Strategy','Head of CI','Product Manager','CMO','Marketing Manager',
                  'Senior Analyst','BI Lead','Chief Strategy Officer','Director of Operations',
                  'VP Sales','Growth Manager','Market Research Lead','Digital Marketing Manager','CEO']
    countries_list = ['United States','United Kingdom','Germany','Canada','Australia','France','India','Singapore']
    deal_stages = ['appointmentscheduled','qualifiedtobuy','presentationscheduled','decisionmakerboughtin','closedwon']

    all_contacts, all_mqls, all_sqls, all_deals = [], [], [], []

    for m, s, lbl in weeks:
        for j in range(np.random.randint(35, 55)):
            day = m + timedelta(days=np.random.randint(0, 7))
            fname, lname = np.random.choice(first_names), np.random.choice(last_names)
            src = np.random.choice(sources, p=[0.35,0.18,0.15,0.10,0.08,0.08,0.06])
            all_contacts.append({
                "email": f"{fname.lower()}.{lname.lower()}{j}@{np.random.choice(companies).lower().replace(' ','')}.com",
                "lifecyclestage": np.random.choice(["subscriber","lead","marketingqualifiedlead","salesqualifiedlead"], p=[0.15,0.40,0.28,0.17]),
                "source": src, "source_detail": "",
                "createdate": day.strftime('%Y-%m-%dT%H:%M:%S.000Z'),
                "name": f"{fname} {lname}", "company": np.random.choice(companies),
                "industry": np.random.choice(industries), "jobtitle": np.random.choice(job_titles),
                "country": np.random.choice(countries_list, p=[0.40,0.15,0.10,0.10,0.08,0.07,0.05,0.05]),
                "lead_status": np.random.choice(["NEW","OPEN","IN_PROGRESS"], p=[0.5,0.3,0.2])})

        for j in range(np.random.randint(9, 18)):
            day = m + timedelta(days=np.random.randint(0, 7))
            fname, lname = np.random.choice(first_names), np.random.choice(last_names)
            all_mqls.append({
                "email": f"{fname.lower()}.{lname.lower()}.mql{j}@{np.random.choice(companies).lower().replace(' ','')}.com",
                "source": np.random.choice(sources, p=[0.40,0.15,0.15,0.10,0.08,0.07,0.05]),
                "createdate": day.strftime('%Y-%m-%dT%H:%M:%S.000Z'),
                "name": f"{fname} {lname}", "company": np.random.choice(companies),
                "industry": np.random.choice(industries), "jobtitle": np.random.choice(job_titles),
                "country": np.random.choice(countries_list, p=[0.40,0.15,0.10,0.10,0.08,0.07,0.05,0.05])})

        for j in range(np.random.randint(3, 9)):
            day = m + timedelta(days=np.random.randint(0, 7))
            fname, lname = np.random.choice(first_names), np.random.choice(last_names)
            all_sqls.append({
                "email": f"{fname.lower()}.{lname.lower()}.sql{j}@{np.random.choice(companies).lower().replace(' ','')}.com",
                "source": np.random.choice(sources[:4], p=[0.45,0.20,0.20,0.15]),
                "createdate": day.strftime('%Y-%m-%dT%H:%M:%S.000Z'),
                "name": f"{fname} {lname}", "company": np.random.choice(companies),
                "industry": np.random.choice(industries), "jobtitle": np.random.choice(job_titles),
                "country": np.random.choice(countries_list, p=[0.40,0.15,0.10,0.10,0.08,0.07,0.05,0.05])})

        for j in range(np.random.randint(2, 6)):
            day = m + timedelta(days=np.random.randint(0, 7))
            comp = np.random.choice(companies)
            amt = np.random.choice([5000,10000,15000,25000,50000,75000], p=[0.25,0.25,0.20,0.15,0.10,0.05])
            all_deals.append({
                "dealname": f"{comp} - {np.random.choice(['Enterprise','Pro','Standard','Custom'])} Plan",
                "amount": float(amt), "dealstage": np.random.choice(deal_stages, p=[0.25,0.25,0.20,0.15,0.15]),
                "pipeline": "default", "createdate": day.strftime('%Y-%m-%dT%H:%M:%S.000Z'),
                "closedate": (day + timedelta(days=np.random.randint(14, 60))).strftime('%Y-%m-%dT%H:%M:%S.000Z'),
                "source": np.random.choice(sources[:4], p=[0.45,0.25,0.18,0.12])})

    return pd.DataFrame(all_contacts), pd.DataFrame(all_mqls), pd.DataFrame(all_sqls), pd.DataFrame(all_deals)


def assign_hubspot_weeks(df, weeks, date_col='createdate'):
    if df.empty:
        df['Week_Idx'] = []; df['Week_Label'] = []; return df
    df = df.copy()
    df['_date'] = pd.to_datetime(df[date_col], errors='coerce')
    df['Week_Idx'] = -1; df['Week_Label'] = None
    for i, (m, s, lbl) in enumerate(weeks):
        mask = (df['_date'].dt.date >= m) & (df['_date'].dt.date <= s)
        df.loc[mask, 'Week_Idx'] = i; df.loc[mask, 'Week_Label'] = lbl
    return df[df['Week_Idx'] >= 0].copy()

# =====================================================
# UI HELPERS
# =====================================================
def fmt(n):
    if pd.isna(n) or n == 0: return "0"
    if abs(n) >= 1_000_000: return f"{n/1e6:.1f}M"
    if abs(n) >= 1_000: return f"{n:,.0f}"
    if isinstance(n, float): return f"{n:.1f}"
    return str(int(n))

def pct_change(cur, prev):
    if prev == 0: return 0
    return ((cur - prev) / prev) * 100

def change_html(val, fmt_str=".1f", suffix="%"):
    cls = "pos" if val >= 0 else "neg"
    arrow = "▲" if val > 0 else ("▼" if val < 0 else "–")
    return f'<span class="{cls}">{arrow} {abs(val):{fmt_str}}{suffix}</span>'

def make_chart(df, x, y, title, chart_type='line', color=None, height=340, labels=True):
    colors = ['#1a56db','#0f3460','#10b981','#f59e0b','#ef4444','#8b5cf6','#ec4899']
    if chart_type == 'line':
        fig = px.line(df, x=x, y=y, title=title, markers=True, color=color, color_discrete_sequence=colors,
                      text=y if (labels and color is None) else None)
        fig.update_traces(line=dict(width=2.5))
        if labels and color is None:
            fig.update_traces(textposition='top center', textfont=dict(size=10, color='#0f3460'), texttemplate='%{text:,.0f}')
    elif chart_type == 'bar':
        fig = px.bar(df, x=x, y=y, title=title, color=color, color_discrete_sequence=colors, barmode='group')
    if color:
        for trace in fig.data:
            trace.hovertemplate = f'<b>{trace.name}</b>: %{{y:,.0f}}<extra></extra>'
    else:
        fig.update_traces(hovertemplate='%{x}<br><b>%{y:,.0f}</b><extra></extra>')
    fig.update_layout(height=height, plot_bgcolor='white', paper_bgcolor='white',
        font=dict(family='Inter', color='#1e293b', size=12),
        margin=dict(t=45, b=40, l=50, r=20),
        xaxis=dict(gridcolor='#f0f0f0', tickfont=dict(size=10), title=None),
        yaxis=dict(gridcolor='#f0f0f0', tickfont=dict(size=10), rangemode='tozero', title=None),
        legend=dict(orientation='h', yanchor='bottom', y=1.02, bgcolor='rgba(0,0,0,0)'),
        hovermode='x unified')
    return fig

def strat_section_start(title, icon=""):
    st.markdown(f'<div class="strat-section"><div class="strat-header">{icon} {title}</div><div class="strat-body">', unsafe_allow_html=True)

def strat_section_end():
    st.markdown('</div></div>', unsafe_allow_html=True)

def strategy_insight(items, style="action"):
    if not items: return
    clean = [i for i in items if i]
    if not clean: return
    bullets = "".join(f"<li>{i}</li>" for i in clean)
    st.markdown(f'<div class="strategy-box {style}"><ul>{bullets}</ul></div>', unsafe_allow_html=True)

def editable_insights(items, key):
    """Editable strategy insights with save/reset."""
    import re
    clean = [re.sub(r'<[^>]+>', '', str(i)).strip() for i in items if i]
    if not clean: return

    edit_key = f'{key}_editing'
    text_key = f'{key}_text'
    area_key = f'{key}_area'
    if edit_key not in st.session_state: st.session_state[edit_key] = False
    auto_text = "\n".join(clean)
    # Always update to latest auto-generated text (unless user is actively editing)
    if text_key not in st.session_state or not st.session_state[edit_key]:
        st.session_state[text_key] = auto_text

    def _toggle(): st.session_state[edit_key] = not st.session_state[edit_key]
    def _save(): st.session_state[text_key] = st.session_state[area_key]; st.session_state[edit_key] = False
    def _reset(): st.session_state[text_key] = auto_text; st.session_state[edit_key] = False

    if st.session_state[edit_key]:
        st.text_area("Edit insights:", value=st.session_state[text_key], height=100, key=area_key)
        c1, c2, c3 = st.columns(3)
        with c1: st.button("Save", use_container_width=True, key=f"{key}_save", on_click=_save)
        with c2: st.button("Reset", use_container_width=True, key=f"{key}_reset", on_click=_reset)
        with c3: st.button("Cancel", use_container_width=True, key=f"{key}_cancel", on_click=_toggle)
    else:
        lines = [l.strip() for l in st.session_state[text_key].split("\n") if l.strip()]
        bullets = "".join(f"<li>{l}</li>" for l in lines)
        st.markdown(f'<div class="strategy-box action"><ul>{bullets}</ul></div>', unsafe_allow_html=True)
        st.button("Edit Insights", key=f"{key}_edit", on_click=_toggle)

# =====================================================
# SIDEBAR
# =====================================================
with st.sidebar:
    st.markdown("### Configuration")
    config = load_config()

    auth_method = st.radio("Data Source", ["demo","oauth"],
        format_func=lambda x: {"demo": "Demo Data", "oauth": "Google OAuth + HubSpot Live"}[x],
        index=["demo","oauth"].index(config.get('auth_method','demo')))
    config['auth_method'] = auth_method; save_config(config)

    if auth_method == "oauth":
        if os.path.exists(TOKEN_FILE):
            st.success("Google signed in")
            if st.button("Sign Out"): os.remove(TOKEN_FILE); st.cache_data.clear(); st.rerun()
        else:
            if st.button("Sign In with Google", type="primary", use_container_width=True):
                creds = get_oauth_credentials()
                if creds: st.rerun()

    st.markdown("---")
    ga4_property = st.text_input("GA4 Property ID", value=config.get('ga4_property','250811000'))
    config['ga4_property'] = ga4_property; save_config(config)

    st.markdown("---")
    st.markdown("**HubSpot**")
    hubspot_token = st.text_input("Private App Token", value=config.get('hubspot_token',''), type="password")
    config['hubspot_token'] = hubspot_token; save_config(config)
    hs_live = bool(hubspot_token and hubspot_token.strip().startswith("pat-"))
    if hs_live:
        st.success("HubSpot connected")
    else:
        st.info("Demo HubSpot data")

    st.markdown("---")
    st.markdown("**Date Range**")
    use_custom = st.checkbox("Custom range", value=False)
    if use_custom:
        custom_end = st.date_input("End Date", value=datetime.now().date())
        num_weeks = st.number_input("Weeks", value=5, min_value=2, max_value=12)
    else:
        custom_end = None; num_weeks = 5

    st.markdown("---")
    st.markdown("**Targets (Weekly)**")
    targets = config.get('targets', DEFAULT_CONFIG['targets'])
    targets['overall_traffic'] = st.number_input("Sessions Target", value=targets.get('overall_traffic', 5538), step=100)
    targets['leads'] = st.number_input("Leads Target", value=targets.get('leads', 50), step=5)
    targets['mqls'] = st.number_input("MQLs Target", value=targets.get('mqls', 13), step=1)
    targets['pipeline'] = st.number_input("Pipeline Target ($)", value=targets.get('pipeline', 50000), step=5000)
    config['targets'] = targets; save_config(config)

    st.markdown("---")
    if st.button("Refresh Data", use_container_width=True, type="primary"):
        st.cache_data.clear(); st.rerun()

# =====================================================
# DATA LOADING
# =====================================================
weeks = get_last_n_weeks(num_weeks, ref_date=custom_end)
cur_idx = len(weeks) - 1
prev_idx = len(weeks) - 2
full_start = weeks[0][0].strftime('%Y-%m-%d')
full_end = weeks[-1][1].strftime('%Y-%m-%d')

with st.spinner("Loading..."):
    if auth_method == "demo":
        d_traffic, d_channels, d_pages, d_organic = gen_demo_ga4(weeks)
    else:
        ct = os.path.getmtime(TOKEN_FILE) if os.path.exists(TOKEN_FILE) else 0
        try:
            d_traffic = fetch_traffic(ct, ga4_property, full_start, full_end)
            d_channels = fetch_channels(ct, ga4_property, full_start, full_end)
            d_pages = fetch_landing_pages(ct, ga4_property, full_start, full_end)
            d_organic = fetch_organic_traffic(ct, ga4_property, full_start, full_end)
        except Exception as e:
            st.error(f"GA4 Error: {e}"); st.stop()

    if hs_live:
        try:
            hs_contacts = fetch_hubspot_contacts(hubspot_token, weeks[0][0], weeks[-1][1])
            hs_deals = fetch_hubspot_deals(hubspot_token, weeks[0][0], weeks[-1][1])
        except Exception as e:
            st.warning(f"HubSpot error: {e}"); hs_contacts, hs_mqls, hs_sqls, hs_deals = gen_demo_hubspot(weeks)
    else:
        hs_contacts, hs_mqls, hs_sqls, hs_deals = gen_demo_hubspot(weeks)

# Assign weeks
t = assign_weeks(d_traffic, weeks)
ch = assign_weeks(d_channels, weeks)
pg = assign_weeks(d_pages, weeks)
og = assign_weeks(d_organic, weeks)
# Filter out OFFLINE source leads globally (case-insensitive, strip whitespace)
def _is_offline(s):
    return str(s).strip().upper() == 'OFFLINE' if pd.notna(s) else False
hs_contacts = hs_contacts[~hs_contacts['source'].apply(_is_offline)] if not hs_contacts.empty else hs_contacts
hs_mqls = derive_mqls_from_contacts(hs_contacts)
hs_sqls = derive_sqls_from_contacts(hs_contacts)

hs_c = assign_hubspot_weeks(hs_contacts, weeks)
hs_m = assign_hubspot_weeks(hs_mqls, weeks)
hs_s = assign_hubspot_weeks(hs_sqls, weeks)
hs_d = assign_hubspot_weeks(hs_deals, weeks)

# Weekly aggs — Sessions from channel sum (matches breakdown), Users from traffic query (avoids double-counting)
ch_wk = ch.groupby(['Week_Idx','Week_Label','Channel']).agg({'Sessions':'sum','Users':'sum'}).reset_index()
_ch_sessions = ch_wk.groupby(['Week_Idx','Week_Label']).agg({'Sessions':'sum'}).reset_index()
_t_users = t.groupby(['Week_Idx','Week_Label']).agg({'Users':'sum','New Users':'sum'}).reset_index()
t_wk = _ch_sessions.merge(_t_users, on=['Week_Idx','Week_Label'], how='left').fillna(0).sort_values('Week_Idx')
pg_wk = pg.groupby(['Week_Idx','Week_Label','Page']).agg({'Sessions':'sum','Users':'sum'}).reset_index()
og_wk = og.groupby(['Week_Idx','Week_Label']).agg({'Sessions':'sum','Users':'sum'}).reset_index().sort_values('Week_Idx')

leads_wk = hs_c.groupby(['Week_Idx','Week_Label']).size().reset_index(name='Leads').sort_values('Week_Idx')
mqls_wk = hs_m.groupby(['Week_Idx','Week_Label']).size().reset_index(name='MQLs').sort_values('Week_Idx')
sqls_wk = hs_s.groupby(['Week_Idx','Week_Label']).size().reset_index(name='SQLs').sort_values('Week_Idx')
deals_wk = hs_d.groupby(['Week_Idx','Week_Label']).agg(Deals=('dealname','count'), Pipeline=('amount','sum'), Avg_Deal=('amount','mean')).reset_index().sort_values('Week_Idx')

# Current week numbers
def _wk_val(df, idx, col):
    r = df[df.Week_Idx==idx]
    return int(r[col].sum()) if len(r) else 0

cur_sessions = _wk_val(t_wk, cur_idx, 'Sessions')
prev_sessions = _wk_val(t_wk, prev_idx, 'Sessions')
cur_users = _wk_val(t_wk, cur_idx, 'Users')
prev_users = _wk_val(t_wk, prev_idx, 'Users')
cur_org = _wk_val(og_wk, cur_idx, 'Sessions')
prev_org = _wk_val(og_wk, prev_idx, 'Sessions')
cur_leads = _wk_val(leads_wk, cur_idx, 'Leads')
prev_leads = _wk_val(leads_wk, prev_idx, 'Leads')
cur_mqls = _wk_val(mqls_wk, cur_idx, 'MQLs')
prev_mqls = _wk_val(mqls_wk, prev_idx, 'MQLs')
cur_sqls = _wk_val(sqls_wk, cur_idx, 'SQLs')
prev_sqls = _wk_val(sqls_wk, prev_idx, 'SQLs')

_d_cur = deals_wk[deals_wk.Week_Idx==cur_idx]
_d_prev = deals_wk[deals_wk.Week_Idx==prev_idx]
cur_deals = int(_d_cur['Deals'].sum()) if len(_d_cur) else 0
prev_deals = int(_d_prev['Deals'].sum()) if len(_d_prev) else 0
cur_pipeline = float(_d_cur['Pipeline'].sum()) if len(_d_cur) else 0
prev_pipeline = float(_d_prev['Pipeline'].sum()) if len(_d_prev) else 0
cur_avg_deal = float(_d_cur['Avg_Deal'].mean()) if len(_d_cur) else 0

# =====================================================
# HEADER
# =====================================================
tag = "DEMO" if auth_method == "demo" else "LIVE"
hs_tag = "HS-LIVE" if hs_live else "HS-DEMO"
st.markdown(f"""
<div class="dash-header">
    <h1>Marketing Strategy Dashboard <span class="tag tag-{'green' if tag=='LIVE' else 'yellow'}">{tag}</span> <span class="tag tag-{'green' if hs_live else 'yellow'}">{hs_tag}</span></h1>
    <p>Week: {weeks[-1][2]} &bull; contify.com &bull; GA4 + HubSpot &bull; {num_weeks}-week view</p>
</div>
""", unsafe_allow_html=True)

# =====================================================
# FUNNEL BAR — Sessions → Leads → MQLs → SQLs → Deals
# =====================================================
def _funnel_card(label, val, prev, conv_from=None, is_last=False):
    chg = pct_change(val, prev)
    color = "#059669" if chg >= 0 else "#dc2626"
    arrow = "▲" if chg > 0 else ("▼" if chg < 0 else "–")
    conv_html = ""
    if conv_from is not None and conv_from > 0:
        rate = val / conv_from * 100
        conv_html = f'<div class="f-conv">{rate:.1f}% conv.</div>'
    arrow_div = "" if is_last else '<div class="f-arrow">→</div>'
    return f'''<div class="funnel-card">
        <div class="f-label">{label}</div>
        <div class="f-val">{fmt(val)}</div>
        <div class="f-change" style="color:{color}">{arrow} {abs(chg):.1f}%</div>
        {conv_html}{arrow_div}
    </div>'''

st.markdown(f'''<div class="funnel-row">
    {_funnel_card("Sessions", cur_sessions, prev_sessions)}
    {_funnel_card("Leads", cur_leads, prev_leads, cur_sessions)}
    {_funnel_card("MQLs", cur_mqls, prev_mqls, cur_leads)}
    {_funnel_card("SQLs", cur_sqls, prev_sqls, cur_mqls)}
    {_funnel_card("Deals", cur_deals, prev_deals, cur_sqls, is_last=True)}
</div>''', unsafe_allow_html=True)

# Pipeline value card
st.markdown(f'''<div style="display:flex;gap:14px;margin-bottom:20px;">
    <div class="m-card" style="flex:1;"><div class="m-label">Pipeline Created</div><div class="m-val">${fmt(cur_pipeline)}</div><div class="m-sub" style="color:{'#059669' if cur_pipeline>=targets['pipeline'] else '#dc2626'}">Target: ${fmt(targets['pipeline'])}</div></div>
    <div class="m-card" style="flex:1;"><div class="m-label">Avg Deal Size</div><div class="m-val">${fmt(cur_avg_deal)}</div><div class="m-sub">{cur_deals} deals this week</div></div>
    <div class="m-card" style="flex:1;"><div class="m-label">Organic Sessions</div><div class="m-val">{fmt(cur_org)}</div><div class="m-sub">{change_html(pct_change(cur_org, prev_org))} vs last week</div></div>
    <div class="m-card" style="flex:1;"><div class="m-label">Overall Conv. Rate</div><div class="m-val">{(cur_deals/cur_sessions*100) if cur_sessions else 0:.2f}%</div><div class="m-sub">Sessions → Deals</div></div>
</div>''', unsafe_allow_html=True)

# =====================================================
# SECTION 1: TRAFFIC OVERVIEW (from GA4)
# =====================================================
strat_section_start("Where is our traffic coming from?", "📊")

c1, c2 = st.columns([3, 2])
with c1:
    fig = make_chart(t_wk, 'Week_Label', 'Sessions', 'Weekly Sessions Trend', height=300)
    st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': False})
with c2:
    # Channel breakdown — current week
    cur_ch = ch_wk[ch_wk.Week_Idx==cur_idx].groupby('Channel')['Sessions'].sum().reset_index().sort_values('Sessions', ascending=False)
    prev_ch = ch_wk[ch_wk.Week_Idx==prev_idx].groupby('Channel')['Sessions'].sum().reset_index()
    ch_merged = cur_ch.merge(prev_ch, on='Channel', suffixes=('','_prev'), how='left').fillna(0)
    ch_merged['Change'] = ch_merged.apply(lambda r: pct_change(r['Sessions'], r['Sessions_prev']), axis=1)

    rows = ""
    for _, r in ch_merged.iterrows():
        rows += f"<tr><td><strong>{r['Channel']}</strong></td><td>{fmt(r['Sessions'])}</td><td>{change_html(r['Change'])}</td></tr>"
    st.markdown(f'<table class="data-table"><tr><th>Channel</th><th>Sessions</th><th>WoW</th></tr>{rows}</table>', unsafe_allow_html=True)

# Top landing pages
with st.expander("Top Landing Pages — This Week"):
    cur_pages = pg_wk[pg_wk.Week_Idx==cur_idx].groupby('Page').agg({'Sessions':'sum','Users':'sum'}).reset_index().sort_values('Sessions', ascending=False).head(10)
    if len(cur_pages):
        total_pg = cur_pages['Sessions'].sum()
        cur_pages['Share'] = (cur_pages['Sessions'] / total_pg * 100).round(1)
        rows = ""
        for i, (_, r) in enumerate(cur_pages.iterrows(), 1):
            rows += f"<tr><td>{i}</td><td><strong>{r['Page']}</strong></td><td>{fmt(r['Sessions'])}</td><td>{fmt(r['Users'])}</td><td>{r['Share']:.1f}%</td></tr>"
        st.markdown(f'<table class="detail-table"><tr><th>#</th><th>Page</th><th>Sessions</th><th>Users</th><th>Share</th></tr>{rows}</table>', unsafe_allow_html=True)

best_ch = ch_merged.iloc[0] if len(ch_merged) else None
strategy_insight([
    f'<strong>{fmt(cur_sessions)}</strong> total sessions this week ({change_html(pct_change(cur_sessions, prev_sessions))} WoW)',
    f'Top channel: <strong>{best_ch["Channel"]}</strong> with {fmt(best_ch["Sessions"])} sessions' if best_ch is not None else '',
    f'Organic contributes <strong>{(cur_org/cur_sessions*100):.0f}%</strong> of total traffic' if cur_sessions > 0 else '',
])
strat_section_end()

# =====================================================
# SECTION 2: LEADS — Who came in?
# =====================================================
strat_section_start("Who are our new leads?", "👥")

_leads_cur = hs_c[hs_c.Week_Idx==cur_idx].copy() if not hs_c.empty else pd.DataFrame()

# Leads by source breakdown first
if not _leads_cur.empty:
    src_counts = _leads_cur.groupby('source').size().reset_index(name='Count').sort_values('Count', ascending=False)
else:
    src_counts = pd.DataFrame(columns=['source', 'Count'])

c1, c2 = st.columns([2, 3])
with c1:
    lead_target_pct = (cur_leads / targets['leads'] * 100) if targets['leads'] > 0 else 0
    pct_color = "#059669" if lead_target_pct >= 100 else ("#f59e0b" if lead_target_pct >= 80 else "#dc2626")
    st.markdown(f'''<div class="m-card" style="margin-bottom:12px;">
        <div class="m-label">Leads This Week</div>
        <div class="m-val">{fmt(cur_leads)}</div>
        <div class="m-sub">{change_html(pct_change(cur_leads, prev_leads))} WoW</div>
        <div style="margin-top:8px;font-size:0.82rem;font-weight:700;color:{pct_color}">{lead_target_pct:.0f}% of target ({fmt(targets['leads'])})</div>
    </div>''', unsafe_allow_html=True)

    # Leads by source mini table
    if not src_counts.empty:
        rows = ""
        for _, r in src_counts.head(6).iterrows():
            rows += f"<tr><td>{r['source'] or 'Unknown'}</td><td><strong>{r['Count']}</strong></td></tr>"
        st.markdown(f'<table class="data-table"><tr><th>Source</th><th>Leads</th></tr>{rows}</table>', unsafe_allow_html=True)

with c2:
    # Leads trend chart
    fig = make_chart(leads_wk, 'Week_Label', 'Leads', 'Leads per Week', height=280)
    st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': False})

# Checkboxes for Total Leads view
show_total_leads = st.checkbox("Show Total Leads Details", value=False, key="chk_total_leads")
if show_total_leads and not _leads_cur.empty:
    src_filter = st.selectbox("Filter by source", ['All'] + sorted(_leads_cur['source'].dropna().unique().tolist()), key="lead_src")
    filtered = _leads_cur if src_filter == 'All' else _leads_cur[_leads_cur.source == src_filter]
    rows = ""
    for _, r in filtered.head(50).iterrows():
        dt = pd.to_datetime(r.get('createdate',''), errors='coerce')
        dt_s = dt.strftime('%d %b') if pd.notna(dt) else '—'
        stage_tag = 'tag-green' if r.get('lifecyclestage','') in ['marketingqualifiedlead','salesqualifiedlead'] else 'tag-blue'
        rows += f"<tr><td><strong>{r.get('name','—')}</strong></td><td>{r.get('company','—')}</td><td>{r.get('industry','—')}</td><td>{r.get('source','—')}</td><td>{r.get('jobtitle','—')}</td><td>{dt_s}</td><td><span class='tag {stage_tag}'>{r.get('lifecyclestage','—')}</span></td></tr>"
    st.markdown(f'<table class="detail-table"><tr><th>Name</th><th>Company</th><th>Industry</th><th>Source</th><th>Title</th><th>Date</th><th>Stage</th></tr>{rows}</table>', unsafe_allow_html=True)
    if len(filtered) > 50: st.caption(f"Showing 50 of {len(filtered)}")

editable_insights([
    f'Total {cur_leads} leads this week ({change_html(pct_change(cur_leads, prev_leads))} WoW) — {lead_target_pct:.0f}% of target',
    f'Top source: {src_counts.iloc[0]["source"]} ({src_counts.iloc[0]["Count"]} leads)' if not src_counts.empty else '',
], key="leads")
strat_section_end()

# =====================================================
# SECTION 3: MQLs — Which leads are converting?
# =====================================================
strat_section_start("Which leads became MQLs?", "🔥")

_mqls_cur = hs_m[hs_m.Week_Idx==cur_idx].copy() if not hs_m.empty else pd.DataFrame()
lead_to_mql = (cur_mqls / cur_leads * 100) if cur_leads > 0 else 0
mql_target_pct = (cur_mqls / targets['mqls'] * 100) if targets['mqls'] > 0 else 0

c1, c2, c3 = st.columns(3)
with c1:
    mql_color = "#059669" if mql_target_pct >= 100 else ("#f59e0b" if mql_target_pct >= 80 else "#dc2626")
    st.markdown(f'<div class="m-card"><div class="m-label">MQLs</div><div class="m-val">{fmt(cur_mqls)}</div><div class="m-sub">{change_html(pct_change(cur_mqls, prev_mqls))} WoW</div><div style="margin-top:6px;font-size:0.8rem;font-weight:700;color:{mql_color}">{mql_target_pct:.0f}% of target</div></div>', unsafe_allow_html=True)
with c2:
    st.markdown(f'<div class="m-card"><div class="m-label">Lead → MQL Rate</div><div class="m-val">{lead_to_mql:.1f}%</div><div class="m-sub">{fmt(cur_mqls)} of {fmt(cur_leads)} leads</div></div>', unsafe_allow_html=True)
with c3:
    # Best converting source
    if not _mqls_cur.empty:
        mql_src = _mqls_cur.groupby('source').size().reset_index(name='MQLs').sort_values('MQLs', ascending=False)
        top_mql_src = mql_src.iloc[0] if len(mql_src) else None
        st.markdown(f'<div class="m-card"><div class="m-label">Top MQL Source</div><div class="m-val">{top_mql_src["source"] if top_mql_src is not None else "—"}</div><div class="m-sub">{fmt(top_mql_src["MQLs"]) if top_mql_src is not None else "0"} MQLs</div></div>', unsafe_allow_html=True)
    else:
        st.markdown('<div class="m-card"><div class="m-label">Top MQL Source</div><div class="m-val">—</div></div>', unsafe_allow_html=True)

# MQL source breakdown + industry
if not _mqls_cur.empty:
    c1, c2 = st.columns(2)
    with c1:
        st.markdown("**By Source**")
        mql_src = _mqls_cur.groupby('source').size().reset_index(name='MQLs').sort_values('MQLs', ascending=False)
        total_m = mql_src['MQLs'].sum()
        rows = ""
        for _, r in mql_src.iterrows():
            share = (r['MQLs']/total_m*100) if total_m else 0
            rows += f"<tr><td><strong>{r['source'] or 'Unknown'}</strong></td><td>{r['MQLs']}</td><td>{share:.0f}%</td></tr>"
        st.markdown(f'<table class="data-table"><tr><th>Source</th><th>MQLs</th><th>Share</th></tr>{rows}</table>', unsafe_allow_html=True)
    with c2:
        st.markdown("**By Industry**")
        mql_ind = _mqls_cur.groupby('industry').size().reset_index(name='MQLs').sort_values('MQLs', ascending=False)
        rows = ""
        for _, r in mql_ind.head(6).iterrows():
            rows += f"<tr><td><strong>{r['industry'] or 'Unknown'}</strong></td><td>{r['MQLs']}</td></tr>"
        st.markdown(f'<table class="data-table"><tr><th>Industry</th><th>MQLs</th></tr>{rows}</table>', unsafe_allow_html=True)

    show_mqls = st.checkbox(f"Show all {len(_mqls_cur)} MQL details", value=False, key="chk_mqls")
    if show_mqls:
        rows = ""
        for _, r in _mqls_cur.head(50).iterrows():
            dt = pd.to_datetime(r.get('createdate',''), errors='coerce')
            dt_s = dt.strftime('%d %b') if pd.notna(dt) else '—'
            rows += f"<tr><td><strong>{r.get('name','—')}</strong></td><td>{r.get('company','—')}</td><td>{r.get('industry','—')}</td><td>{r.get('source','—')}</td><td>{r.get('jobtitle','—')}</td><td>{r.get('country','—')}</td><td>{dt_s}</td></tr>"
        st.markdown(f'<table class="detail-table"><tr><th>Name</th><th>Company</th><th>Industry</th><th>Source</th><th>Title</th><th>Country</th><th>Date</th></tr>{rows}</table>', unsafe_allow_html=True)

editable_insights([
    f'{cur_mqls} MQLs ({change_html(pct_change(cur_mqls, prev_mqls))} WoW) — {mql_target_pct:.0f}% of target',
    f'Lead-to-MQL conversion: {lead_to_mql:.1f}%',
    f'Best source: {mql_src.iloc[0]["source"]} ({mql_src.iloc[0]["MQLs"]} MQLs)' if not _mqls_cur.empty and len(mql_src) else '',
    f'Top industry: {mql_ind.iloc[0]["industry"]} ({mql_ind.iloc[0]["MQLs"]} MQLs)' if not _mqls_cur.empty and len(mql_ind) else '',
], key="mqls")
strat_section_end()

# =====================================================
# SECTION 4: SQLs & Pipeline — What's moving to sales?
# =====================================================
strat_section_start("What moved to sales & pipeline?", "💰")

_sqls_cur = hs_s[hs_s.Week_Idx==cur_idx].copy() if not hs_s.empty else pd.DataFrame()
_deals_cur = hs_d[hs_d.Week_Idx==cur_idx].copy() if not hs_d.empty else pd.DataFrame()
mql_to_sql = (cur_sqls / cur_mqls * 100) if cur_mqls > 0 else 0
sql_to_deal = (cur_deals / cur_sqls * 100) if cur_sqls > 0 else 0

c1, c2, c3, c4 = st.columns(4)
with c1:
    st.markdown(f'<div class="m-card"><div class="m-label">SQLs</div><div class="m-val">{fmt(cur_sqls)}</div><div class="m-sub">{change_html(pct_change(cur_sqls, prev_sqls))} WoW</div></div>', unsafe_allow_html=True)
with c2:
    st.markdown(f'<div class="m-card"><div class="m-label">MQL → SQL</div><div class="m-val">{mql_to_sql:.1f}%</div><div class="m-sub">{fmt(cur_sqls)} of {fmt(cur_mqls)}</div></div>', unsafe_allow_html=True)
with c3:
    pip_color = "#059669" if cur_pipeline >= targets['pipeline'] else "#dc2626"
    st.markdown(f'<div class="m-card"><div class="m-label">Pipeline</div><div class="m-val" style="color:{pip_color}">${fmt(cur_pipeline)}</div><div class="m-sub">Target: ${fmt(targets["pipeline"])}</div></div>', unsafe_allow_html=True)
with c4:
    st.markdown(f'<div class="m-card"><div class="m-label">SQL → Deal</div><div class="m-val">{sql_to_deal:.1f}%</div><div class="m-sub">{fmt(cur_deals)} deals</div></div>', unsafe_allow_html=True)

# SQL details with checkbox
show_sqls = st.checkbox("Show SQLs Created", value=False, key="chk_sqls")
if show_sqls and not _sqls_cur.empty:
    rows = ""
    for _, r in _sqls_cur.iterrows():
        dt = pd.to_datetime(r.get('createdate',''), errors='coerce')
        dt_s = dt.strftime('%d %b') if pd.notna(dt) else '—'
        rows += f"<tr><td><strong>{r.get('name','—')}</strong></td><td>{r.get('company','—')}</td><td>{r.get('industry','—')}</td><td>{r.get('source','—')}</td><td>{r.get('jobtitle','—')}</td><td>{dt_s}</td></tr>"
    st.markdown(f'<table class="data-table"><tr><th>Name</th><th>Company</th><th>Industry</th><th>Source</th><th>Title</th><th>Date</th></tr>{rows}</table>', unsafe_allow_html=True)

# Pipeline Created with checkbox
show_pipeline = st.checkbox("Show Pipeline Created", value=False, key="chk_pipeline")
if show_pipeline and not _deals_cur.empty:
    rows = ""
    for _, r in _deals_cur.iterrows():
        dt = pd.to_datetime(r.get('createdate',''), errors='coerce')
        dt_s = dt.strftime('%d %b') if pd.notna(dt) else '—'
        stage = str(r.get('dealstage','')).replace('_',' ').title()
        rows += f"<tr><td><strong>{r.get('dealname','—')}</strong></td><td><strong>${fmt(r.get('amount',0))}</strong></td><td>{stage}</td><td>{r.get('source','—')}</td><td>{dt_s}</td></tr>"
    st.markdown(f'<table class="detail-table"><tr><th>Deal</th><th>Amount</th><th>Stage</th><th>Source</th><th>Created</th></tr>{rows}</table>', unsafe_allow_html=True)

# Pipeline by source with checkbox
show_pip_src = st.checkbox("Show Pipeline by Source", value=False, key="chk_pip_src")
if show_pip_src and not _deals_cur.empty:
    deal_by_src = _deals_cur.groupby('source').agg(Deals=('dealname','count'), Pipeline=('amount','sum')).reset_index().sort_values('Pipeline', ascending=False)
    rows = ""
    for _, r in deal_by_src.iterrows():
        rows += f"<tr><td><strong>{r['source'] or 'Unknown'}</strong></td><td>{r['Deals']}</td><td><strong>${fmt(r['Pipeline'])}</strong></td></tr>"
    st.markdown(f'<table class="data-table"><tr><th>Source</th><th>Deals</th><th>Pipeline</th></tr>{rows}</table>', unsafe_allow_html=True)

# Pipeline trend
if not deals_wk.empty:
    fig = make_chart(deals_wk, 'Week_Label', 'Pipeline', 'Pipeline Created per Week ($)', chart_type='bar', height=280)
    st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': False})

editable_insights([
    f'{cur_sqls} SQLs, {cur_deals} deals, ${fmt(cur_pipeline)} pipeline this week',
    f'MQL→SQL: {mql_to_sql:.1f}% | SQL→Deal: {sql_to_deal:.1f}%',
    f'Avg deal size: ${fmt(cur_avg_deal)}',
    f'Pipeline {change_html(pct_change(cur_pipeline, prev_pipeline))} vs last week',
], key="pipeline")
strat_section_end()

# =====================================================
# SECTION 5: SOURCE PERFORMANCE — Full Funnel View
# =====================================================
strat_section_start("Which channels drive the best results?", "🎯")

# Build full-funnel by source
_src_leads = hs_c[hs_c.Week_Idx==cur_idx].groupby('source').size().reset_index(name='Leads') if not hs_c.empty else pd.DataFrame(columns=['source','Leads'])
_src_mqls = hs_m[hs_m.Week_Idx==cur_idx].groupby('source').size().reset_index(name='MQLs') if not hs_m.empty else pd.DataFrame(columns=['source','MQLs'])
_src_sqls = hs_s[hs_s.Week_Idx==cur_idx].groupby('source').size().reset_index(name='SQLs') if not hs_s.empty else pd.DataFrame(columns=['source','SQLs'])
_src_deals = hs_d[hs_d.Week_Idx==cur_idx].groupby('source').agg(Deals=('dealname','count'), Pipeline=('amount','sum')).reset_index() if not hs_d.empty else pd.DataFrame(columns=['source','Deals','Pipeline'])

src_funnel = _src_leads.merge(_src_mqls, on='source', how='outer').merge(_src_sqls, on='source', how='outer').merge(_src_deals, on='source', how='outer').fillna(0)
for c in ['Leads','MQLs','SQLs','Deals']: src_funnel[c] = src_funnel[c].astype(int)
src_funnel['L→MQL'] = src_funnel.apply(lambda r: f"{(r['MQLs']/r['Leads']*100):.0f}%" if r['Leads'] > 0 else "—", axis=1)
src_funnel['Pipeline'] = src_funnel['Pipeline'].astype(float)
src_funnel = src_funnel.sort_values('Leads', ascending=False)

if not src_funnel.empty:
    rows = ""
    for _, r in src_funnel.iterrows():
        # Color code the conversion rate
        conv_val = (r['MQLs']/r['Leads']*100) if r['Leads'] > 0 else 0
        conv_color = "#059669" if conv_val >= 30 else ("#f59e0b" if conv_val >= 15 else "#6b7280")
        rows += f"<tr><td><strong>{r['source'] or 'Unknown'}</strong></td><td>{fmt(r['Leads'])}</td><td>{fmt(r['MQLs'])}</td><td>{fmt(r['SQLs'])}</td><td>{fmt(r['Deals'])}</td><td style='color:{conv_color};font-weight:700'>{r['L→MQL']}</td><td><strong>${fmt(r['Pipeline'])}</strong></td></tr>"
    st.markdown(f'<table class="data-table"><tr><th>Source</th><th>Leads</th><th>MQLs</th><th>SQLs</th><th>Deals</th><th>L→MQL Rate</th><th>Pipeline</th></tr>{rows}</table>', unsafe_allow_html=True)

    # Find best performers
    src_funnel['_conv'] = src_funnel.apply(lambda r: (r['MQLs']/r['Leads']*100) if r['Leads'] >= 3 else 0, axis=1)
    best_volume = src_funnel.iloc[0]
    best_conv = src_funnel[src_funnel['_conv']>0].sort_values('_conv', ascending=False).iloc[0] if len(src_funnel[src_funnel['_conv']>0]) else None
    best_pipeline = src_funnel.sort_values('Pipeline', ascending=False).iloc[0]

    strategy_insight([
        f'<strong>Highest volume:</strong> {best_volume["source"]} — {fmt(best_volume["Leads"])} leads',
        f'<strong>Best conversion:</strong> {best_conv["source"]} — {best_conv["_conv"]:.0f}% Lead→MQL' if best_conv is not None else '',
        f'<strong>Most pipeline:</strong> {best_pipeline["source"]} — ${fmt(best_pipeline["Pipeline"])}',
    ])

strat_section_end()

# =====================================================
# SECTION 6: INDUSTRY INTELLIGENCE
# =====================================================
strat_section_start("Which industries are performing?", "🏭")

_ind_leads = hs_c[hs_c.Week_Idx==cur_idx].groupby('industry').size().reset_index(name='Leads') if not hs_c.empty else pd.DataFrame(columns=['industry','Leads'])
_ind_mqls = hs_m[hs_m.Week_Idx==cur_idx].groupby('industry').size().reset_index(name='MQLs') if not hs_m.empty else pd.DataFrame(columns=['industry','MQLs'])
_ind_sqls = hs_s[hs_s.Week_Idx==cur_idx].groupby('industry').size().reset_index(name='SQLs') if not hs_s.empty else pd.DataFrame(columns=['industry','SQLs'])

ind_perf = _ind_leads.merge(_ind_mqls, on='industry', how='outer').merge(_ind_sqls, on='industry', how='outer').fillna(0)
for c in ['Leads','MQLs','SQLs']: ind_perf[c] = ind_perf[c].astype(int)
ind_perf['MQL Rate'] = ind_perf.apply(lambda r: f"{(r['MQLs']/r['Leads']*100):.0f}%" if r['Leads'] > 0 else "—", axis=1)
ind_perf = ind_perf[ind_perf['industry']!=''].sort_values('Leads', ascending=False)

if not ind_perf.empty:
    c1, c2 = st.columns([3, 2])
    with c1:
        rows = ""
        for _, r in ind_perf.iterrows():
            rows += f"<tr><td><strong>{r['industry']}</strong></td><td>{fmt(r['Leads'])}</td><td>{fmt(r['MQLs'])}</td><td>{fmt(r['SQLs'])}</td><td>{r['MQL Rate']}</td></tr>"
        st.markdown(f'<table class="data-table"><tr><th>Industry</th><th>Leads</th><th>MQLs</th><th>SQLs</th><th>MQL Rate</th></tr>{rows}</table>', unsafe_allow_html=True)
    with c2:
        fig = px.bar(ind_perf.head(6), x='industry', y='MQLs', title='MQLs by Industry', color_discrete_sequence=['#1a56db'])
        fig.update_layout(height=280, plot_bgcolor='white', paper_bgcolor='white', font=dict(family='Inter', size=11),
            margin=dict(t=40,b=60,l=40,r=10), xaxis=dict(title=None, tickangle=-25), yaxis=dict(title=None, rangemode='tozero'))
        fig.update_traces(hovertemplate='%{x}<br><b>%{y}</b> MQLs<extra></extra>')
        st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': False})

    # Strategy recommendation
    ind_perf['_mql_num'] = ind_perf.apply(lambda r: (r['MQLs']/r['Leads']*100) if r['Leads'] >= 2 else 0, axis=1)
    best_ind = ind_perf.sort_values('_mql_num', ascending=False).iloc[0] if len(ind_perf) else None
    strategy_insight([
        f'<strong>{ind_perf.iloc[0]["industry"]}</strong> leads in volume ({fmt(ind_perf.iloc[0]["Leads"])} leads)',
        f'<strong>{best_ind["industry"]}</strong> has the best MQL rate ({best_ind["_mql_num"]:.0f}%) — consider doubling down on content for this vertical' if best_ind is not None and best_ind['_mql_num'] > 0 else '',
    ])

strat_section_end()

# =====================================================
# SECTION 7: WEEKLY TRENDS
# =====================================================
strat_section_start("Week-over-Week Trend", "📈")

# Combined funnel trend
trend_df = t_wk[['Week_Idx','Week_Label','Sessions']].merge(leads_wk, on=['Week_Idx','Week_Label'], how='left')
trend_df = trend_df.merge(mqls_wk, on=['Week_Idx','Week_Label'], how='left')
trend_df = trend_df.merge(sqls_wk, on=['Week_Idx','Week_Label'], how='left')
trend_df = trend_df.merge(deals_wk[['Week_Idx','Week_Label','Deals','Pipeline']], on=['Week_Idx','Week_Label'], how='left').fillna(0)

# Display table — current week first
display_trend = trend_df.sort_values('Week_Idx', ascending=False)
rows = ""
for _, r in display_trend.iterrows():
    current_tag = ' <span class="tag-current">CURRENT</span>' if r['Week_Idx'] == cur_idx else ''
    rows += f"<tr><td><strong>{r['Week_Label']}</strong>{current_tag}</td><td>{fmt(r['Sessions'])}</td><td>{fmt(r['Leads'])}</td><td>{fmt(r['MQLs'])}</td><td>{fmt(r['SQLs'])}</td><td>{fmt(r['Deals'])}</td><td>${fmt(r['Pipeline'])}</td></tr>"
st.markdown(f'<table class="data-table"><tr><th>Week</th><th>Sessions</th><th>Leads</th><th>MQLs</th><th>SQLs</th><th>Deals</th><th>Pipeline</th></tr>{rows}</table>', unsafe_allow_html=True)

# Trend charts
c1, c2 = st.columns(2)
with c1:
    melt_df = trend_df.melt(id_vars=['Week_Label'], value_vars=['Leads','MQLs','SQLs'], var_name='Stage', value_name='Count')
    fig = make_chart(melt_df, 'Week_Label', 'Count', 'Funnel Trend', color='Stage', height=280, labels=False)
    st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': False})
with c2:
    fig = make_chart(trend_df, 'Week_Label', 'Pipeline', 'Pipeline Trend ($)', chart_type='bar', height=280)
    st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': False})

strat_section_end()

# =====================================================
# SECTION 8: STRATEGIC ACTION ITEMS
# =====================================================
strat_section_start("Strategy & Action Items", "🧠")

# Auto-generate strategic recommendations based on data
actions = []

# Traffic health
if pct_change(cur_sessions, prev_sessions) < -5:
    actions.append(f'<span class="neg">Traffic dropped {abs(pct_change(cur_sessions, prev_sessions)):.1f}%</span> — investigate channel performance and content publishing cadence')
elif pct_change(cur_sessions, prev_sessions) > 10:
    actions.append(f'<span class="pos">Traffic up {pct_change(cur_sessions, prev_sessions):.1f}%</span> — analyze which pages drove the spike and replicate')

# Lead quality
if cur_leads > 0:
    if lead_to_mql < 15:
        actions.append(f'Lead→MQL rate is <span class="neg">{lead_to_mql:.1f}%</span> (low) — review lead scoring criteria and landing page qualification')
    elif lead_to_mql > 35:
        actions.append(f'Lead→MQL rate is <span class="pos">{lead_to_mql:.1f}%</span> (excellent) — current targeting is working well')

# Pipeline
if cur_pipeline < targets['pipeline'] * 0.5:
    actions.append(f'Pipeline at <span class="neg">{(cur_pipeline/targets["pipeline"]*100):.0f}%</span> of target — need more SQLs or higher deal values')
elif cur_pipeline >= targets['pipeline']:
    actions.append(f'Pipeline <span class="pos">exceeded target</span> at ${fmt(cur_pipeline)} — focus on deal acceleration and close rates')

# Best channel recommendation
if not src_funnel.empty:
    best_conv_src = src_funnel[src_funnel['_conv']>0].sort_values('_conv', ascending=False)
    if len(best_conv_src):
        top = best_conv_src.iloc[0]
        actions.append(f'<strong>{top["source"]}</strong> has {top["_conv"]:.0f}% MQL conversion — consider increasing investment in this channel')

# Industry focus
if not ind_perf.empty and len(ind_perf):
    top_ind = ind_perf.iloc[0]
    actions.append(f'<strong>{top_ind["industry"]}</strong> is the top vertical — create targeted content and case studies for this industry')

# MQL target
if mql_target_pct < 80:
    actions.append(f'MQLs at <span class="neg">{mql_target_pct:.0f}%</span> of target — review nurture sequences and consider new lead magnets')

if not actions:
    actions = ['All metrics are tracking well this week — maintain current strategy']

editable_insights(actions, key="strategy")

strat_section_end()

# =====================================================
# DOWNLOAD
# =====================================================
st.markdown("---")
st.markdown("### Download Report")
dl1, dl2, dl3 = st.columns(3)

# Prepare report data
import re
def _strip(t): return re.sub(r'<[^>]+>', '', str(t)).strip()

report_data = {
    'week_label': weeks[-1][2],
    'kpi_rows': [
        ('Sessions', fmt(targets['overall_traffic']), fmt(cur_sessions), f"{(cur_sessions/targets['overall_traffic']*100):.0f}%" if targets['overall_traffic'] else '0%'),
        ('Leads', fmt(targets['leads']), fmt(cur_leads), f"{(cur_leads/targets['leads']*100):.0f}%" if targets['leads'] else '0%'),
        ('MQLs', fmt(targets['mqls']), fmt(cur_mqls), f"{(cur_mqls/targets['mqls']*100):.0f}%" if targets['mqls'] else '0%'),
        ('SQLs', fmt(targets.get('sqls',6)), fmt(cur_sqls), f"{(cur_sqls/targets.get('sqls',6)*100):.0f}%" if targets.get('sqls',6) else '0%'),
        ('Pipeline', f"${fmt(targets['pipeline'])}", f"${fmt(cur_pipeline)}", f"{(cur_pipeline/targets['pipeline']*100):.0f}%" if targets['pipeline'] else '0%'),
    ],
    'kpi_insights': [_strip(a) for a in actions],
    'sections': [
        {'title': 'Source Performance', 'table': {'columns': ['Source','Leads','MQLs','SQLs','Deals','L→MQL','Pipeline'], 'data': [[r.get('source',''),fmt(r['Leads']),fmt(r['MQLs']),fmt(r['SQLs']),fmt(r['Deals']),r['L→MQL'],f"${fmt(r['Pipeline'])}"] for _,r in src_funnel.iterrows()]} if not src_funnel.empty else {'columns':[],'data':[]}, 'insights': []},
        {'title': 'Industry Performance', 'table': {'columns': ['Industry','Leads','MQLs','SQLs','MQL Rate'], 'data': [[r['industry'],fmt(r['Leads']),fmt(r['MQLs']),fmt(r['SQLs']),r['MQL Rate']] for _,r in ind_perf.iterrows()]} if not ind_perf.empty else {'columns':[],'data':[]}, 'insights': []},
        {'title': 'Weekly Trend', 'table': {'columns': ['Week','Sessions','Leads','MQLs','SQLs','Deals','Pipeline'], 'data': [[r['Week_Label'],fmt(r['Sessions']),fmt(r['Leads']),fmt(r['MQLs']),fmt(r['SQLs']),fmt(r['Deals']),f"${fmt(r['Pipeline'])}"] for _,r in display_trend.iterrows()]}, 'insights': []},
    ]
}

with dl1:
    try:
        from fpdf import FPDF
        pdf = FPDF(); pdf.set_auto_page_break(auto=True, margin=15); pdf.add_page()
        pdf.set_fill_color(15,52,96); pdf.rect(10,10,190,25,'F')
        pdf.set_text_color(255,255,255); pdf.set_font('Helvetica','B',16); pdf.set_xy(15,14)
        pdf.cell(0,10,'Marketing Strategy Dashboard — Contify',ln=True)
        pdf.set_font('Helvetica','',9); pdf.set_xy(15,24); pdf.cell(0,8,f"Week: {report_data['week_label']}  |  {datetime.now().strftime('%B %d, %Y')}",ln=True)
        pdf.ln(10); pdf.set_text_color(30,41,59)
        pdf.set_font('Helvetica','B',12); pdf.cell(0,8,'Funnel Summary',ln=True); pdf.ln(2)
        pdf.set_font('Helvetica','B',8); pdf.set_fill_color(15,52,96); pdf.set_text_color(255,255,255)
        for h in ['Metric','Target','Achieved','%']: pdf.cell(40,7,h,1,0,'C',True)
        pdf.ln(); pdf.set_text_color(30,41,59); pdf.set_font('Helvetica','',8)
        for row in report_data['kpi_rows']:
            for v in row: pdf.cell(40,6,str(v)[:20],1,0,'C')
            pdf.ln()
        pdf.ln(4); pdf.set_font('Helvetica','I',8)
        for ins in report_data['kpi_insights']: pdf.multi_cell(0,4,f"  * {ins}")
        for sec in report_data['sections']:
            pdf.ln(4); pdf.set_font('Helvetica','B',11); pdf.cell(0,8,sec['title'],ln=True); pdf.ln(1)
            if sec.get('table') and sec['table'].get('columns'):
                cols = sec['table']['columns']; data = sec['table']['data']
                cw = min(int(180/len(cols)),35)
                pdf.set_font('Helvetica','B',7); pdf.set_fill_color(240,244,255); pdf.set_text_color(15,52,96)
                for c in cols: pdf.cell(cw,6,str(c)[:18],1,0,'C',True)
                pdf.ln(); pdf.set_text_color(30,41,59); pdf.set_font('Helvetica','',7)
                for row in data[:12]:
                    for v in row: pdf.cell(cw,5,str(v)[:18],1,0,'C')
                    pdf.ln()
        st.download_button("Download PDF", pdf.output(), f"Marketing_Strategy_{weeks[-1][0].strftime('%Y%m%d')}.pdf", "application/pdf", use_container_width=True)
    except ImportError: st.warning("Install fpdf2: `pip install fpdf2`")
    except Exception as e: st.error(f"PDF: {e}")

with dl2:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document(); style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(10)
        t_ = doc.add_heading('Marketing Strategy Dashboard — Contify', level=0)
        for run in t_.runs: run.font.color.rgb = RGBColor(15,52,96)
        doc.add_paragraph(f"Week: {report_data['week_label']}  |  {datetime.now().strftime('%B %d, %Y')}")
        doc.add_heading('Funnel Summary', level=1)
        tbl = doc.add_table(rows=1+len(report_data['kpi_rows']), cols=4); tbl.style = 'Light Grid Accent 1'
        for i,h in enumerate(['Metric','Target','Achieved','%']): tbl.rows[0].cells[i].text = h
        for ri,row in enumerate(report_data['kpi_rows']):
            for ci,v in enumerate(row): tbl.rows[ri+1].cells[ci].text = str(v)
        doc.add_heading('Action Items', level=1)
        for ins in report_data['kpi_insights']: doc.add_paragraph(ins, style='List Bullet')
        for sec in report_data['sections']:
            doc.add_heading(sec['title'], level=1)
            if sec.get('table') and sec['table'].get('columns'):
                cols = sec['table']['columns']; data = sec['table']['data']
                tbl = doc.add_table(rows=1+min(len(data),15), cols=len(cols)); tbl.style = 'Light Grid Accent 1'
                for i,c in enumerate(cols): tbl.rows[0].cells[i].text = str(c)
                for ri,row in enumerate(data[:15]):
                    for ci,v in enumerate(row): tbl.rows[ri+1].cells[ci].text = str(v)
        import io; buf = io.BytesIO(); doc.save(buf)
        st.download_button("Download Word", buf.getvalue(), f"Marketing_Strategy_{weeks[-1][0].strftime('%Y%m%d')}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
    except ImportError: st.warning("Install python-docx: `pip install python-docx`")
    except Exception as e: st.error(f"Word: {e}")

with dl3:
    import io
    csv_buf = io.StringIO()
    csv_buf.write("=== WEEKLY FUNNEL ===\n")
    display_trend[['Week_Label','Sessions','Leads','MQLs','SQLs','Deals','Pipeline']].to_csv(csv_buf, index=False)
    if not src_funnel.empty:
        csv_buf.write("\n=== SOURCE PERFORMANCE ===\n")
        src_funnel[['source','Leads','MQLs','SQLs','Deals','L→MQL','Pipeline']].to_csv(csv_buf, index=False)
    if not ind_perf.empty:
        csv_buf.write("\n=== INDUSTRY PERFORMANCE ===\n")
        ind_perf[['industry','Leads','MQLs','SQLs','MQL Rate']].to_csv(csv_buf, index=False)
    if not _leads_cur.empty:
        csv_buf.write("\n=== LEADS DETAIL ===\n")
        exp_cols = [c for c in ['name','company','industry','source','createdate','lifecyclestage','jobtitle','country'] if c in _leads_cur.columns]
        _leads_cur[exp_cols].to_csv(csv_buf, index=False)
    st.download_button("Download CSV", csv_buf.getvalue(), f"Marketing_Strategy_{weeks[-1][0].strftime('%Y%m%d')}.csv", "text/csv", use_container_width=True)

# Footer
st.markdown(f'<div style="text-align:center;padding:24px 0 8px;color:#9ca3af;font-size:0.78rem;">{datetime.now().strftime("%B %d, %Y at %I:%M %p")} &bull; Contify Marketing Strategy Dashboard &bull; GA4 + HubSpot</div>', unsafe_allow_html=True)
