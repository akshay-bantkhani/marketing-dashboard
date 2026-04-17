"""HTML → DOCX conversion — DocxBuilder class and recursive element processor."""

import io
from urllib.parse import urljoin

from bs4 import NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

from .config import (
    CAPTION_FONT_SIZE_PT,
    DEFAULT_FONT_NAME,
    DEFAULT_FONT_SIZE_PT,
    IMAGE_DPI,
    MAX_IMAGE_WIDTH_INCHES,
)
from .scraper import download_image


# ---------------------------------------------------------------------------
# Hyperlink helper
# ---------------------------------------------------------------------------


def add_hyperlink(paragraph, url: str, text: str):
    """Add a clickable hyperlink to a DOCX paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = paragraph._element.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})

    run = paragraph._element.makeelement(qn("w:r"), {})
    rPr = paragraph._element.makeelement(qn("w:rPr"), {})

    color = paragraph._element.makeelement(qn("w:color"), {qn("w:val"): "0563C1"})
    underline = paragraph._element.makeelement(qn("w:u"), {qn("w:val"): "single"})
    rPr.append(color)
    rPr.append(underline)
    run.append(rPr)

    t = paragraph._element.makeelement(qn("w:t"), {})
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)

    hyperlink.append(run)
    paragraph._element.append(hyperlink)


# ---------------------------------------------------------------------------
# DocxBuilder
# ---------------------------------------------------------------------------


class DocxBuilder:
    """Incrementally builds a DOCX document from parsed HTML content."""

    def __init__(self, base_url: str):
        self.doc = Document()
        self.base_url = base_url
        self._set_default_font()

    # -- setup --------------------------------------------------------------

    def _set_default_font(self):
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = DEFAULT_FONT_NAME
        font.size = Pt(DEFAULT_FONT_SIZE_PT)

    def _resolve_url(self, url: str) -> str:
        """Turn a relative URL into an absolute one."""
        if not url:
            return ""
        if url.startswith("//"):
            return "https:" + url
        if url.startswith(("http://", "https://")):
            return url
        return urljoin(self.base_url, url)

    # -- block-level elements -----------------------------------------------

    def add_heading(self, text: str, level: int):
        docx_level = min(level, 4)
        self.doc.add_heading(text, level=docx_level)

    def add_paragraph_text(self, text: str):
        text = text.strip()
        if text:
            self.doc.add_paragraph(text)

    def add_image(self, img_url: str, alt_text: str = ""):
        """Download and embed an image with optional caption."""
        resolved = self._resolve_url(img_url)
        if not resolved:
            return
        result = download_image(resolved)
        if not result:
            return

        data, _ext = result
        stream = io.BytesIO(data)
        try:
            img = Image.open(io.BytesIO(data))
            w, _h = img.size
            width_inches = min(w / IMAGE_DPI, MAX_IMAGE_WIDTH_INCHES)

            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(stream, width=Inches(width_inches))
        except Exception:
            return

        if alt_text:
            caption = self.doc.add_paragraph(alt_text)
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.style = self.doc.styles["Normal"]
            for run in caption.runs:
                run.italic = True
                run.font.size = Pt(CAPTION_FONT_SIZE_PT)
                run.font.color.rgb = RGBColor(128, 128, 128)

    def add_list_item(self, text: str, ordered: bool = False, level: int = 0):
        style = "List Number" if ordered else "List Bullet"
        p = self.doc.add_paragraph(text, style=style)
        if level > 0:
            p.paragraph_format.left_indent = Inches(0.5 * level)

    def add_table(self, rows: list[list[str]]):
        if not rows:
            return
        num_cols = max(len(r) for r in rows)
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.style = "Table Grid"
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                if j < num_cols:
                    table.rows[i].cells[j].text = cell_text
        # Bold header row
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    def add_blockquote(self, text: str):
        para = self.doc.add_paragraph(text)
        para.paragraph_format.left_indent = Inches(0.5)
        for run in para.runs:
            run.italic = True

    # -- inline / rich content ---------------------------------------------

    def add_rich_paragraph(self, element: Tag):
        """Add a paragraph with mixed content (text + links + bold/italic)."""
        para = self.doc.add_paragraph()
        self._process_inline(element, para)

    def _process_inline(self, element, paragraph, bold=False, italic=False):
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text.strip():
                    run = paragraph.add_run(text)
                    run.bold = bold
                    run.italic = italic
            elif isinstance(child, Tag):
                if child.name == "a":
                    href = child.get("href", "")
                    link_text = child.get_text()
                    if href and link_text.strip():
                        add_hyperlink(paragraph, self._resolve_url(href), link_text)
                    elif link_text.strip():
                        paragraph.add_run(link_text)
                elif child.name in ("strong", "b"):
                    self._process_inline(child, paragraph, bold=True, italic=italic)
                elif child.name in ("em", "i"):
                    self._process_inline(child, paragraph, bold=bold, italic=True)
                elif child.name == "br":
                    paragraph.add_run("\n")
                else:
                    self._process_inline(child, paragraph, bold=bold, italic=italic)

    # -- save ---------------------------------------------------------------

    def save(self, filepath: str):
        self.doc.save(filepath)


# ---------------------------------------------------------------------------
# Recursive HTML → DOCX element processor
# ---------------------------------------------------------------------------


def _get_image_src(el: Tag) -> str:
    """Extract the best available image source from an <img> tag."""
    src = (
        el.get("data-src")
        or el.get("data-lazy-src")
        or el.get("data-orig-file")
        or el.get("src")
        or ""
    )
    if src.startswith("data:"):
        src = el.get("data-src") or el.get("data-lazy-src") or ""
    return src


def process_element(el: Tag, builder: DocxBuilder):
    """Recursively walk an HTML element tree and add content to *builder*."""
    if isinstance(el, NavigableString):
        text = str(el).strip()
        if text:
            builder.add_paragraph_text(text)
        return
    if not isinstance(el, Tag):
        return

    tag = el.name

    # -- Headings -----------------------------------------------------------
    if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
        text = el.get_text(strip=True)
        if text:
            builder.add_heading(text, level=int(tag[1]) - 1)
        return

    # -- Images -------------------------------------------------------------
    if tag == "img":
        src = _get_image_src(el)
        if src:
            builder.add_image(src, el.get("alt", ""))
        return

    # -- Figure (image + caption) -------------------------------------------
    if tag == "figure":
        img = el.find("img")
        if img:
            src = _get_image_src(img)
            figcaption = el.find("figcaption")
            caption = figcaption.get_text(strip=True) if figcaption else img.get("alt", "")
            if src:
                builder.add_image(src, caption)
        return

    # -- Tables -------------------------------------------------------------
    if tag == "table":
        rows = []
        for tr in el.find_all("tr"):
            cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
            if cells:
                rows.append(cells)
        if rows:
            builder.add_table(rows)
        return

    # -- Lists --------------------------------------------------------------
    if tag in ("ul", "ol"):
        ordered = tag == "ol"
        for li in el.find_all("li", recursive=False):
            parts = []
            for child in li.children:
                if isinstance(child, NavigableString):
                    parts.append(str(child).strip())
                elif isinstance(child, Tag) and child.name not in ("ul", "ol"):
                    parts.append(child.get_text(strip=True))
            text = " ".join(p for p in parts if p)
            if text:
                builder.add_list_item(text, ordered=ordered)
            # Nested lists
            for nested in li.find_all(["ul", "ol"], recursive=False):
                nested_ordered = nested.name == "ol"
                for nested_li in nested.find_all("li", recursive=False):
                    nested_text = nested_li.get_text(strip=True)
                    if nested_text:
                        builder.add_list_item(nested_text, ordered=nested_ordered, level=1)
        return

    # -- Paragraphs ---------------------------------------------------------
    if tag == "p":
        img = el.find("img")
        if img:
            src = _get_image_src(img)
            if src:
                builder.add_image(src, img.get("alt", ""))

        if el.find(["a", "strong", "b", "em", "i", "span"]):
            builder.add_rich_paragraph(el)
        else:
            text = el.get_text(strip=True)
            if text:
                builder.add_paragraph_text(text)
        return

    # -- Blockquote ---------------------------------------------------------
    if tag == "blockquote":
        text = el.get_text(strip=True)
        if text:
            builder.add_blockquote(text)
        return

    # -- Container elements — recurse into children -------------------------
    if tag in ("div", "section", "article", "main", "span", "aside"):
        for child in el.children:
            if isinstance(child, Tag):
                process_element(child, builder)
            elif isinstance(child, NavigableString):
                text = str(child).strip()
                if text and len(text) > 20:
                    builder.add_paragraph_text(text)
        return

    # -- Default: recurse ---------------------------------------------------
    for child in el.children:
        if isinstance(child, Tag):
            process_element(child, builder)
